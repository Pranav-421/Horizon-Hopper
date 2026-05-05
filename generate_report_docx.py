"""Generate Horizon Hopper Project Report as .docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

PRIMARY = RGBColor(0, 104, 95)  # #00685F
DARK = RGBColor(10, 25, 49)
GRAY = RGBColor(100, 100, 100)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
    return h


def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return t


def para(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p


# ═══════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
para("SRM Institute of Science and Technology", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY)
para("Department of Networking and Communications", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
doc.add_paragraph()
para("21CSE210P — Agentic AI \"NOVA\"", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY)
doc.add_paragraph()
doc.add_paragraph()
para("🌏 HORIZON HOPPER", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY)
para("AI-Powered Agentic Travel Planner for Tamil Nadu", italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY)
doc.add_paragraph()
doc.add_paragraph()
para("Version 3.0.0", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Date: May 2026", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ═══════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════
heading("Table of Contents", level=1)
toc_items = [
    "1. Problem Statement",
    "2. Agentic AI Model — Pipeline Diagram Prompt",
    "3. Model Design",
    "4. Agentic View — Architecture Table",
    "5. Technologies Used",
    "6. Decision Making Strategies",
    "7. Challenges / Limitations",
    "8. Future Scope",
    "9. Project Structure",
    "10. API Endpoints",
    "11. Database Schema",
    "12. Dataset Coverage",
    "13. Running the Application",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")
doc.add_page_break()

# ═══════════════════════════════════════════
#  1. PROBLEM STATEMENT
# ═══════════════════════════════════════════
heading("1. Problem Statement", level=1)

p = doc.add_paragraph()
r = p.add_run("Problem Statement: ")
r.bold = True
r.font.color.rgb = PRIMARY
r.font.size = Pt(12)
r2 = p.add_run(
    "Design an Agentic AI system that generates personalized, data-driven travel itineraries "
    "across Tamil Nadu by orchestrating a pipeline of specialized agents — each responsible for "
    "intent detection, commute planning, hotel recommendation, attraction curation, and itinerary "
    "synthesis — to replace static template-based planners with an intelligent, adaptive travel assistant."
)
r2.italic = True

heading("1.1 Problem Context", level=2)
para(
    "Traditional travel planners rely on rigid templates and manual search, failing to adapt to "
    "individual preferences like transport mode, dietary needs, budget constraints, or trip purpose. "
    "Travelers in Tamil Nadu face a fragmented information landscape: railway timetables, metro routes, "
    "hotel listings, and tourist attractions are spread across dozens of platforms with no unified, intelligent interface."
)

heading("1.2 Proposed Solution", level=2)
para(
    "Horizon Hopper is a multi-agent AI travel planner that employs a pipeline of 8 specialized agents — "
    "orchestrated by a central coordinator — to ingest a natural-language travel request and produce a "
    "complete, personalized travel plan in a single API call. The system covers 18+ cities across Tamil Nadu "
    "with CSV-grounded datasets ensuring factual accuracy, while leveraging the Groq-hosted Llama 3.3 70B "
    "model for natural-language reasoning and refinement."
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  2. PIPELINE DIAGRAM PROMPT
# ═══════════════════════════════════════════
heading("2. Agentic AI Model — Pipeline Diagram Prompt", level=1)
para("Use the following prompt to generate the Agentic AI pipeline/workflow diagram:", bold=True, size=11)

prompt_text = (
    "Create a professional, clean architectural pipeline diagram for an Agentic AI travel planner "
    "called 'Horizon Hopper'. The diagram should flow left-to-right (or top-to-bottom) and show the following:\n\n"
    "1. User Input (a person icon or form) sends a natural-language travel request.\n"
    "2. Orchestrator (central hub/coordinator box) receives the request and manages the pipeline.\n"
    "3. Perception Layer - three boxes in parallel: Intent Agent (classifies trip purpose), "
    "Location Agent (geo-resolves source/destination), and Memory Agent (loads user preferences from SQLite).\n"
    "4. Decision Layer - three boxes: Commute Agent (generates transport options using LLM + static data), "
    "Stay Agent (recommends hotels from CSV), and Attraction Agent (curates POIs from CSV + Wikipedia GeoSearch).\n"
    "5. Action Layer - two boxes: Itinerary Agent (synthesizes the final plan via Groq LLM), "
    "and Feedback Agent (refines specific sections based on user feedback).\n"
    "6. Infrastructure Layer - shown as a horizontal foundation bar running beneath all agent layers, "
    "containing four components: Config Service (pydantic-settings, .env loading), Security Module "
    "(JWT token creation/verification + bcrypt password hashing), SQLite Database (WAL-mode, schema "
    "auto-migration, users/memory/trips tables), and Formatter Service (regex parsers converting raw "
    "agent text into structured Pydantic JSON). Draw bidirectional arrows from this layer up to each "
    "agent layer showing that all agents depend on config, security gates the API, database persists "
    "state, and formatter transforms final output.\n"
    "7. Output - flows from the Formatter Service to the Next.js Frontend displaying the results.\n\n"
    "Use a modern, minimal design style with rounded rectangles, subtle gradients (teal/dark green primary "
    "color #00685F), connecting arrows showing data flow, and small icons inside each agent box. Include a "
    "sidebar showing external data sources: CSV Datasets, Groq API, OSRM API, Wikipedia GeoSearch. "
    "Add a legend for Perception/Decision/Action/Infrastructure color coding (use grey or blue for Infrastructure). "
    "The overall look should be professional enough for an academic presentation at SRM University."
)
p = doc.add_paragraph()
r = p.add_run(prompt_text)
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GRAY
doc.add_page_break()

# ═══════════════════════════════════════════
#  3. MODEL DESIGN
# ═══════════════════════════════════════════
heading("3. Model Design", level=1)

heading("3.1 Application Details", level=2)
para(
    "Horizon Hopper is a full-stack agentic AI web application consisting of a FastAPI v0.115 backend "
    "serving a REST API with JWT-authenticated endpoints, a Next.js 15 frontend with a premium glassmorphism UI, "
    "dark mode, and interactive Leaflet maps, an SQLite (WAL-mode) database for persistent user profiles, "
    "and 8 specialized AI agents orchestrated through a central pipeline coordinator."
)

heading("3.2 Smartness / Novelty", level=2)
add_table(
    ["Aspect", "Implementation"],
    [
        ["Hybrid AI + Rule-Based", "Every agent uses LLM (Groq Llama 3.3 70B) as primary, with deterministic rule-based fallback — ensuring 100% uptime"],
        ["CSV-Grounded Generation", "Attractions, hotels, and train schedules are sourced from curated CSV datasets first, then enriched by LLM — minimizing hallucination"],
        ["Persistent Memory", "The Memory Agent learns user preferences across sessions and injects them into every future query"],
        ["Targeted Refinement", "Feedback Agent detects which section feedback targets and re-generates only that section"],
        ["Real-Time Geo-Enrichment", "Falls back to Wikipedia GeoSearch API for live attraction discovery when destination isn't in bundled CSV"],
    ],
)

heading("3.3 Target Users", level=2)
add_table(
    ["Persona", "Use Case"],
    [
        ["Office Commuters", "Daily route optimization across Chennai metro, MTC bus, and ride-hailing"],
        ["Business Travelers", "Multi-day trip planning with premium hotel selection"],
        ["Leisure Tourists", "Weekend getaway planning with attraction curation and budget-aware itineraries"],
        ["Student Backpackers", "Ultra-budget travel with hostel-only stays and cheapest transport modes"],
    ],
)

heading("3.4 Agentic AI Technique", level=2)
para(
    "The system implements a sequential agent pipeline pattern (inspired by LangGraph workflows):"
)
steps = [
    "Orchestrator receives the user request and assembles a TravelState dictionary",
    "Agents execute in fixed order: Intent → Location → Memory → Commute → Stay → Attraction → Itinerary",
    "Each agent reads upstream results and writes its output back to the shared state",
    "The Formatter service parses raw agent text outputs into structured Pydantic models for the frontend",
    "On refinement, the Feedback Agent identifies the affected section and selectively re-invokes only the relevant upstream agent",
]
for s in steps:
    doc.add_paragraph(s, style="List Number")
doc.add_page_break()

# ═══════════════════════════════════════════
#  4. AGENTIC VIEW TABLE
# ═══════════════════════════════════════════
heading("4. Agentic View — Architecture Table", level=1)
add_table(
    ["Sl.No", "Agentic Architecture", "Application/Model", "Decision Making Strategies", "Challenges/Limitation"],
    [
        [
            "1",
            "Perception Module",
            "Intent Agent (intent_agent.py)\nLocation Agent (maps_tool.py)\nMemory Agent (memory_agent.py)",
            "• Keyword rule matching for 4 intent labels\n• Groq LLM zero-temp classification\n• Fuzzy name matching + alias resolution\n• Haversine distance computation\n• SQLite-backed preference persistence",
            "• Limited to 4 intent categories\n• No multi-intent support\n• 23 pre-indexed cities only\n• Keyword-only preference extraction",
        ],
        [
            "2",
            "Decision Module",
            "Commute Agent (commute_agent.py)\nStay Agent (stay_agent.py)\nAttraction Agent (attraction_agent.py)",
            "• Distance-based transport mode eligibility matrix\n• Real train schedule CSV lookup\n• Chennai metro route computation\n• LLM enhancement for descriptions\n• Multi-factor hotel scoring (budget + intent + rating)\n• CSV-first attraction lookup, Wikipedia GeoSearch fallback",
            "• No real-time traffic data\n• Static CSV train schedules\n• Distance-based cab pricing estimates\n• Static hotel data (no live availability)\n• Wikipedia may return non-tourist entries",
        ],
        [
            "3",
            "Action Module",
            "Itinerary Agent (itinerary_agent.py)\nFeedback Agent (feedback_agent.py)",
            "• Groq LLM (Llama 3.3 70B, temp 0.7) generates narrative itinerary\n• System prompt constrains to verified attractions\n• Seasonal weather advisory injected\n• Keyword-based section detection for targeted refinement\n• Static template fallback when LLM unavailable",
            "• LLM output quality varies\n• 800 token limit may truncate\n• No multi-day scheduling logic\n• Section detection may misclassify\n• Cannot handle cross-section feedback",
        ],
        [
            "4",
            "Infrastructure Layer",
            "Config Service (core/config.py)\nSecurity Module (core/security.py)\nDatabase Layer (db/database.py)\nFormatter Service (services/formatter.py)",
            "• Centralized pydantic-settings with .env auto-loading\n• JWT token creation (HS256, 24hr expiry) + bcrypt hashing\n• SQLite WAL-mode with auto-schema + legacy JSON migration\n• Context-managed connections (auto-commit/rollback)\n• Regex-based parsers for commute, stay, attraction text\n• Map data builder with Haversine bounds\n• Assembles PlannerResponse (22 Pydantic models)",
            "• Single .env file (no per-environment profiles)\n• No key rotation or refresh tokens\n• SQLite single-writer limitation\n• No ORM (raw SQL)\n• Regex parsers brittle to LLM format changes",
        ],
    ],
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  5. TECHNOLOGIES USED
# ═══════════════════════════════════════════
heading("5. Technologies Used", level=1)

heading("5.1 Models", level=2)
add_table(
    ["Model", "Provider", "Purpose", "Configuration"],
    [["Llama 3.3 70B Versatile", "Groq Cloud API", "Intent classification, commute enhancement, itinerary generation", "Temp: 0-0.7; Max tokens: 16-800"]],
)

heading("5.2 Frameworks", level=2)
add_table(
    ["Framework", "Version", "Role"],
    [
        ["FastAPI", "≥ 0.115", "Backend REST API server with OpenAPI docs"],
        ["Next.js", "15", "Frontend React framework with App Router & SSR"],
        ["React", "19", "Component-based UI with hooks-driven state"],
        ["Pydantic", "≥ 2.8", "Request/response validation (22 schema models)"],
        ["pydantic-settings", "≥ 2.0", "Environment-driven configuration"],
        ["Pandas", "≥ 2.0", "CSV dataset loading and processing"],
        ["Leaflet", "—", "Interactive map rendering with route polylines"],
    ],
)

heading("5.3 Deployment Tools", level=2)
add_table(
    ["Tool", "Purpose"],
    [
        ["Uvicorn", "ASGI server for FastAPI (with hot-reload)"],
        ["SQLite (WAL mode)", "Embedded DB for users, memory, trips"],
        ["PyJWT", "JWT token creation/verification (HS256, 24hr expiry)"],
        ["bcrypt", "Password hashing with legacy fallback"],
        ["OSRM API", "Open-source road routing geometry"],
        ["Wikipedia GeoSearch API", "Live attraction discovery for unlisted cities"],
        ["python-dotenv", "Environment variable loading from .env"],
    ],
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  6. DECISION MAKING STRATEGIES
# ═══════════════════════════════════════════
heading("6. Decision Making Strategies", level=1)

heading("6.1 Intent Classification (Utility Function)", level=2)
para("The Intent Agent employs a two-tier classification strategy:")
para("Priority 1: Groq LLM (zero-temperature) → strict label output")
para("Priority 2: Rule-based keyword matching → deterministic fallback")
para(
    "LLM Path: Sends the user query with a system prompt constraining output to one of four labels. "
    "Temperature = 0 ensures deterministic classification."
)
para(
    "Rule Path: Scans for keywords (e.g., 'temple' → leisure_trip; 'meeting' → business_trip) "
    "with priority ordering that checks job_interview before leisure_trip to avoid false positives."
)

heading("6.2 Transport Scoring (Commute Agent)", level=2)
add_table(
    ["Transport Mode", "Eligibility Condition", "Cost Formula"],
    [
        ["Metro", "Both endpoints have metro stations", "max(20, distance × 2.2)"],
        ["MTC Bus", "Both in Chennai AND distance < 30km", "max(10, distance × 1.2)"],
        ["TNSTC/SETC", "Distance > 30km (inter-city)", "max(80, distance × 1.5)"],
        ["Train", "Distance > 30km + corridor exists", "CSV lookup or max(30, distance × 0.9)"],
        ["Cab/Auto", "Any distance", "max(100, distance × 18)"],
        ["Bike Taxi", "Distance < 25km", "max(40, distance × 8)"],
        ["Self-Drive", "Distance > 100km", "max(2000, distance × 12)"],
    ],
)

heading("6.3 Hotel Recommendation (Stay Agent)", level=2)
para("The Stay Agent uses a multi-factor scoring model:")
para("• Budget fit: +3 (within budget), +1 (within budget + ₹1000)")
para("• Intent-category alignment: +2 (e.g., business → premium/luxury)")
para("• Explicit preference match: +3 (e.g., user says 'luxury')")
para("• Tiebreak: rating (descending)")
doc.add_page_break()

# ═══════════════════════════════════════════
#  7. CHALLENGES
# ═══════════════════════════════════════════
heading("7. Challenges / Limitations", level=1)
add_table(
    ["Category", "Limitation", "Impact"],
    [
        ["Data Freshness", "CSV datasets are static snapshots", "Stale pricing or schedule discrepancies"],
        ["Geographic Scope", "Limited to 18 TN cities with pre-indexed data", "Reduced quality for lesser-known destinations"],
        ["LLM Dependency", "Groq API rate limits and latency", "Potential degradation under high load"],
        ["No Real-Time Data", "No live traffic, weather, or availability", "Plans based on heuristic estimates"],
        ["Single-Intent", "Classifies to exactly one label per request", "Cannot handle compound queries"],
        ["No Booking", "System recommends but cannot book", "Users must book manually"],
        ["Memory Extraction", "Keyword-based, not NLU-powered", "May miss nuanced preferences"],
    ],
)
doc.add_page_break()

# ═══════════════════════════════════════════
#  8. FUTURE SCOPE
# ═══════════════════════════════════════════
heading("8. Future Scope", level=1)

heading("8.1 Short-Term (Next 3 Months)", level=2)
items_short = [
    "Live API Integration — IRCTC for real-time trains, Google Places for live hotel pricing, OpenWeatherMap for weather",
    "Multi-Intent Support — LLM-based multi-label classification with confidence scores",
    "Voice Input — Web Speech API for hands-free trip planning on mobile",
    "Automated Testing — pytest suite with mock Groq responses",
]
for item in items_short:
    doc.add_paragraph(item, style="List Bullet")

heading("8.2 Medium-Term (3–6 Months)", level=2)
items_mid = [
    "Booking Integration — MakeMyTrip/Goibibo APIs for one-click booking",
    "Multi-State Expansion — Karnataka, Kerala, Andhra Pradesh coverage",
    "Group Travel — Budget splitting, shared itineraries, preference conflict resolution",
    "Offline Mode — PWA with cached itineraries for offline access",
]
for item in items_mid:
    doc.add_paragraph(item, style="List Bullet")

heading("8.3 Long-Term (6–12 Months)", level=2)
items_long = [
    "Reinforcement Learning from Feedback — Fine-tune custom model from real traveler behavior",
    "Multi-Modal Agent Communication — LangGraph state graphs for sophisticated workflows",
    "AR Navigation — Augmented reality wayfinding at attractions",
    "Carbon Footprint Tracker — Environmental impact display per transport option",
]
for item in items_long:
    doc.add_paragraph(item, style="List Bullet")
doc.add_page_break()

# ═══════════════════════════════════════════
#  9. PROJECT STRUCTURE
# ═══════════════════════════════════════════
heading("9. Project Structure", level=1)
structure = """horizon-hopper/
├── frontend/                     # Next.js 15 + React 19
│   ├── src/app/                  # App Router pages + API proxy
│   ├── src/components/           # 5 React components
│   ├── src/lib/                  # Session, proxy, types
│   └── public/assets/            # 270+ city & stay images
│
├── backend/                      # FastAPI 0.115
│   ├── main.py                   # Entry point, CORS, error handler
│   ├── app/agents/               # 8 AI agents + orchestrator
│   ├── app/api/routes.py         # 7 REST endpoints (JWT auth)
│   ├── app/core/                 # Config + Security
│   ├── app/db/database.py        # SQLite layer (CRUD, migration)
│   ├── app/models/schemas.py     # 22 Pydantic schema models
│   ├── app/services/             # Orchestrator + Formatter
│   ├── app/utils/                # 6 data tools
│   └── data/                     # 5 CSV datasets
│
└── README.md"""
p = doc.add_paragraph()
r = p.add_run(structure)
r.font.size = Pt(9)
r.font.name = "Consolas"

# ═══════════════════════════════════════════
#  10. API ENDPOINTS
# ═══════════════════════════════════════════
heading("10. API Endpoints", level=1)
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/api/auth/login", "—", "Authenticate user, return JWT + profile"],
        ["GET", "/api/users", "—", "List all user profiles with memory"],
        ["GET", "/api/users/{id}/memory", "—", "Get user preferences & trip history"],
        ["POST", "/api/users/{id}/feedback", "—", "Save feedback to user memory"],
        ["POST", "/api/trips/plan", "JWT", "Generate complete travel plan"],
        ["POST", "/api/trips/refine", "—", "Refine specific section via feedback"],
        ["GET", "/api/graph/data", "—", "Serve knowledge graph JSON"],
    ],
)

# ═══════════════════════════════════════════
#  11. DATABASE SCHEMA
# ═══════════════════════════════════════════
heading("11. Database Schema", level=1)
schema_text = """CREATE TABLE users (
    id            TEXT PRIMARY KEY,
    name          TEXT NOT NULL,
    avatar        TEXT DEFAULT 'U',
    password_hash TEXT NOT NULL,
    created_at    TEXT DEFAULT (datetime('now'))
);

CREATE TABLE memory (
    user_id              TEXT PRIMARY KEY REFERENCES users(id),
    preferred_transport  TEXT,
    food_preference      TEXT,
    budget_range         TEXT,
    accommodation_type   TEXT,
    avoid                TEXT DEFAULT '[]',
    feedback_notes       TEXT DEFAULT '[]',
    updated_at           TEXT DEFAULT (datetime('now'))
);

CREATE TABLE trips (
    id          INTEGER PRIMARY KEY AUTOINCREMENT,
    user_id     TEXT REFERENCES users(id),
    date        TEXT NOT NULL,
    intent      TEXT DEFAULT 'leisure_trip',
    summary     TEXT DEFAULT '',
    destination TEXT DEFAULT ''
);"""
p = doc.add_paragraph()
r = p.add_run(schema_text)
r.font.size = Pt(9)
r.font.name = "Consolas"

# ═══════════════════════════════════════════
#  12. DATASET COVERAGE
# ═══════════════════════════════════════════
heading("12. Dataset Coverage", level=1)
add_table(
    ["Dataset", "Records", "Coverage"],
    [
        ["attractions.csv", "88 attractions", "18 cities, 6 categories"],
        ["stays.csv", "90+ hotels", "5 tiers (Luxury → Hostel)"],
        ["railway_schedule.csv", "30+ routes", "Inter-city train connections"],
        ["metro_routes.csv", "40+ stations", "Chennai metro network"],
        ["offices.csv", "25+ locations", "IT parks & business centers"],
    ],
)

# ═══════════════════════════════════════════
#  13. RUNNING THE APP
# ═══════════════════════════════════════════
heading("13. Running the Application", level=1)
para("Prerequisites: Python 3.12+, Node.js 18+, Groq API Key (free tier)", bold=True)

heading("Demo Accounts", level=2)
add_table(
    ["Username", "Password", "Profile"],
    [
        ["arjun", "arjun123", "Business traveler, vegetarian"],
        ["priya", "priya456", "Leisure explorer, budget-conscious"],
        ["vikram", "vikram789", "Student backpacker"],
    ],
)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Report generated May 2026 | Horizon Hopper v3.0.0 | Built with ❤ for Tamil Nadu travelers")
r.font.size = Pt(9)
r.font.color.rgb = GRAY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), "Horizon_Hopper_Report_v2.docx")
doc.save(output_path)
print(f"[OK] Report saved to: {output_path}")
